"""
Génère le script vidéo de présentation Digital School + NTT S.A.R.L.

Trois durées :
  - 3 minutes  (teaser / réseaux sociaux)
  - 10 minutes (présentation commerciale)
  - 25 minutes (tour complet des fonctionnalités)

Usage :
    python docs/generer_script_video.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_FILE = OUT_DIR / "Script_Video_Digital_School_NTT.docx"


def _font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def h0(doc, text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _font(run, bold=bold, italic=italic)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    _font(run)


def vo(doc, text):
    """Voix off."""
    para = doc.add_paragraph()
    lab = para.add_run("VOIX OFF : ")
    _font(lab, bold=True, color=RGBColor(0, 119, 197))
    body = para.add_run(text)
    _font(body, italic=True)


def ecran(doc, text):
    para = doc.add_paragraph()
    lab = para.add_run("ÉCRAN : ")
    _font(lab, bold=True, color=RGBColor(22, 101, 52))
    body = para.add_run(text)
    _font(body)


def plan(doc, text):
    para = doc.add_paragraph()
    lab = para.add_run("PLAN / DÉCOUPAGE : ")
    _font(lab, bold=True, color=RGBColor(146, 64, 14))
    body = para.add_run(text)
    _font(body)


def timed(doc, start, end, titre):
    para = doc.add_paragraph()
    lab = para.add_run(f"[{start} → {end}] ")
    _font(lab, bold=True, color=RGBColor(120, 53, 15))
    body = para.add_run(titre)
    _font(body, bold=True)


def add_ntt_block(doc):
    """Bloc présentation NTT S.A.R.L (commun aux 3 versions, à adapter)."""
    h(doc, "Présentation de NTT S.A.R.L (à placer en ouverture)", 2)
    p(
        doc,
        "Complétez les crochets avec vos infos officielles avant le tournage.",
        italic=True,
    )
    bullet(doc, "Raison sociale : NTT S.A.R.L")
    bullet(doc, "Nature : startup technologique / éditeur de solutions numériques")
    bullet(doc, "Produit phare : plateforme Digital School")
    bullet(doc, "Siège / zone d’intervention : [Ville, pays — ex. Kinshasa, RDC]")
    bullet(doc, "Année de création : [AAAA]")
    bullet(doc, "Mission : digitaliser la gestion des établissements scolaires")
    bullet(
        doc,
        "Positionnement : solutions concrètes pour le terrain éducatif africain "
        "(multi-écoles, finances CDF/USD, WhatsApp, pédagogie, portails)",
    )
    bullet(doc, "Contact / site / réseaux : [e-mail] · [téléphone] · [URL / LinkedIn]")
    bullet(doc, "Équipe à l’écran (optionnel) : [noms + rôles]")


def build_common_intro(doc):
    h0(doc, "Script vidéo — Digital School par NTT S.A.R.L")
    p(
        doc,
        "Trois versions de montage pour présenter toutes les fonctionnalités de "
        "l’application Digital School, avec une séquence dédiée à NTT S.A.R.L, "
        "la startup conceptrice.",
        italic=True,
    )
    p(
        doc,
        "Format recommandé : Full HD 1080p, 16:9, français français. "
        "Captures d’écran + voix off + logo NTT / Digital School. "
        "Musique douce sous VO, baisser pendant les démos écran.",
        italic=True,
    )

    h(doc, "1. Objectifs du film", 1)
    bullet(doc, "Présenter NTT S.A.R.L comme concepteur de la solution.")
    bullet(doc, "Couvrir l’ensemble des modules de Digital School.")
    bullet(doc, "Disposer de 3 durées selon le canal (réseaux, démo commerciale, formation).")

    h(doc, "2. Versions disponibles", 1)
    bullet(doc, "Version A — 3 minutes : teaser / pitch court")
    bullet(doc, "Version B — 10 minutes : présentation commerciale complète")
    bullet(doc, "Version C — 25 minutes : tour détaillé de toutes les fonctionnalités")

    add_ntt_block(doc)

    h(doc, "3. Habillage graphique commun", 1)
    bullet(doc, "Générique début : logo NTT S.A.R.L → transition → logo Digital School")
    bullet(doc, "Bandeau bas (option) : « Conçu par NTT S.A.R.L »")
    bullet(doc, "Générique fin : contacts + appel à l’action (démo / essai pilote)")
    bullet(doc, "Couleur d’accent : bleu Digital School (#0077C5)")


def build_version_3min(doc):
    h(doc, "VERSION A — 3 minutes (teaser)", 1)
    p(doc, "Usage : LinkedIn, WhatsApp Business, site web, salon, accroche avant démo live.")
    p(doc, "Rythme : dynamique. Une idée = un plan. Pas de formulaire long.")

    timed(doc, "0:00", "0:25", "NTT S.A.R.L")
    ecran(doc, "Logo NTT S.A.R.L, siège / équipe (photo), carte RDC / Afrique si pertinent.")
    vo(
        doc,
        "NTT S.A.R.L est une startup technologique qui conçoit des solutions numériques "
        "pour les organisations africaines. Notre produit phare pour l’éducation : Digital School.",
    )

    timed(doc, "0:25", "0:45", "Le problème")
    ecran(doc, "Montage rapide : cahiers, Excel, messages WhatsApp dispersés, files d’attente à la caisse.")
    vo(
        doc,
        "Dans beaucoup d’écoles, inscriptions, paiements, notes et communication parents "
        "sont encore dispersés. Résultat : perte de temps, litiges, manque de visibilité.",
    )

    timed(doc, "0:45", "1:10", "La solution")
    ecran(doc, "Dashboard Digital School + menu Inscriptions / Finances / Pédagogie / GRH.")
    vo(
        doc,
        "Digital School regroupe toute la gestion scolaire dans une plateforme unique, "
        "multi-écoles, avec des espaces dédiés pour la direction, la caisse, "
        "les enseignants, les parents et les élèves.",
    )

    timed(doc, "1:10", "2:20", "Fonctionnalités en flash (6 plans × ~12 s)")
    plan(
        doc,
        "1) Inscriptions élèves/tuteurs · 2) Encaissement + reçu · 3) WhatsApp paiement · "
        "4) Notes & bulletins · 5) Visio cours + questions · 6) Portail parent",
    )
    vo(
        doc,
        "Inscriptions et classes. Finances : minerval, taux CDF/USD, reçus et notifications WhatsApp. "
        "Pédagogie : notes, bulletins, présences. Cours en ligne par visioconférence avec questions en direct. "
        "Portails parent, élève et enseignant. Et la GRH pour le personnel et la paie.",
    )

    timed(doc, "2:20", "2:45", "Différenciation")
    ecran(doc, "Split screen : back-office + téléphone parent recevant WhatsApp + salle visio.")
    vo(
        doc,
        "Pensé pour le terrain : année scolaire nationale, multi-devises, communication parents, "
        "et enseignement à distance intégré.",
    )

    timed(doc, "2:45", "3:00", "Appel à l’action")
    ecran(doc, "Logo Dual NTT + Digital School. Texte : Demandez une démo. Contacts.")
    vo(
        doc,
        "Digital School, conçu par NTT S.A.R.L. Demandez une démonstration pour votre établissement.",
    )


def build_version_10min(doc):
    h(doc, "VERSION B — 10 minutes (présentation commerciale)", 1)
    p(doc, "Usage : rendez-vous école, webinar, proposition commerciale, YouTube.")
    p(doc, "Rythme : clair, démonstratif. Montrer l’écran réel de l’app.")

    timed(doc, "0:00", "1:00", "NTT S.A.R.L")
    ecran(doc, "Générique NTT → pitch startup (mission, produit, terrain).")
    vo(
        doc,
        "Bienvenue. Je vous présente Digital School, conçu par NTT S.A.R.L, "
        "startup spécialisée dans les solutions numériques pour l’éducation. "
        "Notre ambition : donner à chaque école une colonne vertébrale digitale — "
        "fiable, simple et adaptée aux réalités locales.",
    )
    plan(doc, "Carton texte : Mission · Produit · Pour qui (écoles privées, publiques, conventionnées).")

    timed(doc, "1:00", "1:40", "Promesse produit")
    ecran(doc, "Page connexion + dashboard direction.")
    vo(
        doc,
        "Digital School centralise inscriptions, finances, pédagogie, ressources humaines "
        "et communication. Chaque acteur se connecte avec son rôle : direction, caissier, "
        "enseignant, parent ou élève.",
    )

    timed(doc, "1:40", "3:00", "Inscriptions & structure")
    ecran(doc, "Classes → fiche élève → tuteur (popup) → inscription année en cours.")
    vo(
        doc,
        "Vous organisez vos classes, créez les fiches élèves et parents, "
        "et inscrivez pour l’année scolaire en cours. Les matricules sont générés automatiquement. "
        "La capacité des classes est respectée.",
    )

    timed(doc, "3:00", "5:00", "Finances")
    ecran(
        doc,
        "Frais scolaires (Minerval) → encaissement → reçu → taux de change → WhatsApp → aperçu compta.",
    )
    vo(
        doc,
        "Côté finances : barèmes de frais, encaissements multi-devises CDF et USD, "
        "taux de change enregistré, reçus, demandes de correction tracées. "
        "À chaque paiement validé, le parent peut être notifié sur WhatsApp. "
        "La comptabilité suit : plan comptable, journaux, grand livre, balance.",
    )

    timed(doc, "5:00", "6:30", "Pédagogie")
    ecran(doc, "Travaux cotés → notes → bulletin → appel de présence.")
    vo(
        doc,
        "Côté pédagogie, Digital School suit le modèle des bulletins RDC : "
        "travaux journaliers, examens, périodes. Les enseignants saisissent les notes ; "
        "le titulaire gère l’appel et les bulletins.",
    )

    timed(doc, "6:30", "8:00", "Cours en ligne & visio")
    ecran(doc, "Planifier une séance → démarrer → salle Jitsi → panneau Questions.")
    vo(
        doc,
        "Les cours peuvent aussi se dispenser à distance : visioconférence intégrée, "
        "les élèves rejoignent depuis leur portail, et posent des questions en direct. "
        "L’enseignant répond, épingle ou traite oralement. "
        "Des ressources e-learning complètent la visio.",
    )

    timed(doc, "8:00", "9:10", "Portails & GRH")
    ecran(doc, "Portail parent / élève / enseignant → GRH personnel & paie.")
    vo(
        doc,
        "Parents : suivi des enfants, frais, notes, annonces, messagerie. "
        "Élèves : notes et cours en ligne. Enseignants : cockpit de classe. "
        "Et la GRH pour le personnel, contrats, congés, présences et paie.",
    )

    timed(doc, "9:10", "10:00", "Closing NTT + CTA")
    ecran(doc, "Récap icônes modules + logos NTT / Digital School + contacts.")
    vo(
        doc,
        "Digital School, conçu et développé par NTT S.A.R.L. "
        "Une seule plateforme pour piloter votre établissement. "
        "Contactez-nous pour une démonstration ou un essai pilote.",
    )


def build_version_25min(doc):
    h(doc, "VERSION C — 25 minutes (tour complet des fonctionnalités)", 1)
    p(
        doc,
        "Usage : formation, onboarding école, vidéo YouTube « produit en détail », "
        "support commercial approfondi.",
    )
    p(doc, "Rythme : pédagogique. Chapitrage à l’écran obligatoire.")

    # Table of chapters
    h(doc, "Chapitrage recommandé (timecodes)", 2)
    chapters = [
        ("0:00–2:00", "NTT S.A.R.L & générique"),
        ("2:00–3:30", "Vue d’ensemble & rôles"),
        ("3:30–6:30", "Inscriptions, classes, élèves, tuteurs"),
        ("6:30–11:00", "Finances complètes"),
        ("11:00–14:30", "Pédagogie : matières, notes, bulletins, présences"),
        ("14:30–18:00", "Cours en ligne, visio, questions"),
        ("18:00–21:00", "Portails parent / élève / enseignant"),
        ("21:00–23:00", "GRH & salaires"),
        ("23:00–25:00", "Synthèse, NTT, appel à l’action"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Timecode"
    table.rows[0].cells[1].text = "Chapitre"
    for a, b in chapters:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    timed(doc, "0:00", "2:00", "NTT S.A.R.L")
    ecran(
        doc,
        "Logo NTT → présentation société (mission, équipe, vision) → logo Digital School.",
    )
    vo(
        doc,
        "NTT S.A.R.L est la startup qui a conçu Digital School. "
        "Nous développons des outils numériques concrets pour moderniser la gestion scolaire : "
        "administration, finances, pédagogie et relation avec les familles. "
        "Digital School est le fruit de cette vision.",
    )
    plan(doc, "Cartons : Qui sommes-nous · Pourquoi l’éducation · Notre produit.")

    timed(doc, "2:00", "3:30", "Vue d’ensemble & rôles")
    ecran(doc, "Connexion → dashboard → tour du menu latéral.")
    vo(
        doc,
        "Digital School est multi-établissements. Chaque école a son espace isolé. "
        "Les rôles définissent l’accès : manager, directeur, trésorier, caissier, "
        "enseignant, parent, élève.",
    )

    timed(doc, "3:30", "6:30", "Inscriptions")
    ecran(
        doc,
        "Synthèse inscriptions → classes & salles → fiche élève → tuteur (liste globale + popup) "
        "→ nouvelle inscription → badge déjà inscrit → communications direction.",
    )
    vo(
        doc,
        "Module Inscriptions : synthèse des effectifs, gestion des classes avec capacité et titulaire, "
        "fiches élèves avec photo et nationalité, parents et tuteurs avec résidence, "
        "inscription sur l’année scolaire en cours, communications de la direction vers les parents.",
    )

    timed(doc, "6:30", "11:00", "Finances — toutes les fonctions")
    ecran(
        doc,
        "Dashboard finances → frais scolaires (Minerval, Ajouter type) → paiements élèves → "
        "encaisser → demandes de correction → taux de change → WhatsApp → "
        "payer le personnel → plan comptable → journaux → écritures → grand livre → balance.",
    )
    vo(
        doc,
        "Module Finances : définition des frais scolaires et du minerval, encaissement des paiements "
        "élèves avec modes de paiement, reçus, suivi des soldes. "
        "Demandes de correction pour sécuriser la caisse. "
        "Taux de change CDF/USD. Notifications WhatsApp automatiques aux parents. "
        "Paiement des salaires du personnel. "
        "Comptabilité : plan comptable, journaux, écritures, grand livre et balance.",
    )

    timed(doc, "11:00", "14:30", "Pédagogie")
    ecran(
        doc,
        "Matières → périodes bulletin → affectations enseignants → travaux cotés → notes → "
        "bulletins classe/élève → évolution → présences & récap.",
    )
    vo(
        doc,
        "Module Pédagogie : catalogue de matières, périodes de bulletin, affectation des enseignants, "
        "création de travaux cotés et saisie des notes, bulletins conformes au modèle RDC, "
        "suivi d’évolution des élèves, appel de présence réservé au titulaire, récapitulatifs.",
    )

    timed(doc, "14:30", "18:00", "Cours en ligne & visioconférence")
    ecran(
        doc,
        "Cours en ligne enseignant : planifier séance → démarrer → salle visio Jitsi → "
        "questions (poser, répondre, épingler, traité oralement) → côté élève rejoindre. "
        "Puis aperçu ressources e-learning (chapitres / leçons).",
    )
    vo(
        doc,
        "Dispensation des cours en ligne par visioconférence : l’enseignant planifie, démarre et anime "
        "la séance ; les élèves rejoignent depuis leur espace. "
        "Un système de questions permet l’interaction pendant le cours. "
        "En complément, des cours enregistrés structurés en chapitres et leçons, "
        "avec suivi de progression.",
    )

    timed(doc, "18:00", "21:00", "Portails")
    ecran(
        doc,
        "Enseignant : dashboard classes, travaux, messages, présences, cours. "
        "Élève : notes, cours en ligne. "
        "Parent : enfants, frais, notes, présences, annonces, messagerie.",
    )
    vo(
        doc,
        "Trois portails « Mon espace ». L’enseignant pilote ses classes. "
        "L’élève consulte ses notes et suit les cours. "
        "Le parent suit ses enfants, les paiements, les annonces et échange avec l’école.",
    )

    timed(doc, "21:00", "23:00", "GRH")
    ecran(doc, "Personnel → contrats → congés → présences → paies → lien Finances salaires.")
    vo(
        doc,
        "Module GRH : fiches personnel, contrats, demandes de congés, pointage des présences, "
        "génération et suivi des paies, paiement des salaires depuis les finances.",
    )

    timed(doc, "23:00", "25:00", "Synthèse & closing NTT")
    ecran(
        doc,
        "Mosaique des modules → « Conçu par NTT S.A.R.L » → contacts → CTA démo / pilote.",
    )
    vo(
        doc,
        "Voilà Digital School : inscriptions, finances, pédagogie, cours en visio, "
        "portails et GRH — une plateforme complète pour votre établissement. "
        "Conçue par NTT S.A.R.L. Contactez-nous pour une démonstration personnalisée "
        "ou un essai pilote.",
    )


def build_production_notes(doc):
    h(doc, "4. Notes de production", 1)
    h(doc, "4.1 Assets à préparer", 2)
    bullet(doc, "Logo NTT S.A.R.L (PNG transparent) + charte couleurs")
    bullet(doc, "Logo Digital School (déjà dans static/images/logo-ds.png)")
    bullet(doc, "Compte démo avec données réalistes (classe, élèves, paiement, séance visio)")
    bullet(doc, "Captures / enregistrement écran (OBS) en 1080p")
    bullet(doc, "Musique libre de droits, bas volume sous VO")
    bullet(doc, "Sous-titres FR recommandés (accessibilité + LinkedIn)")

    h(doc, "4.2 Règles de tournage écran", 2)
    bullet(doc, "Curseur visible, clics lents, zoom ponctuel sur les champs importants")
    bullet(doc, "Masquer données personnelles réelles (utiliser jeu de démo)")
    bullet(doc, "Couper les temps de chargement")
    bullet(doc, "Pour la visio : planifier une courte séance et montrer le panneau Questions")

    h(doc, "4.3 Dérivés marketing", 2)
    bullet(doc, "Extraire 5 à 8 clips de 15–20 s depuis la version 10 ou 25 min")
    bullet(doc, "Version 3 min = master réseaux")
    bullet(doc, "Version 10 min = page produit / proposition commerciale")
    bullet(doc, "Version 25 min = formation / YouTube chapitreé")

    h(doc, "5. Texte générique de fin (identique aux 3 versions)", 1)
    vo(
        doc,
        "Digital School — conçu par NTT S.A.R.L. "
        "Modernisez la gestion de votre école. Demandez une démo.",
    )
    ecran(
        doc,
        "Digital School · NTT S.A.R.L · [téléphone] · [e-mail] · [site / LinkedIn]",
    )


def build():
    doc = Document()
    build_common_intro(doc)
    build_version_3min(doc)
    build_version_10min(doc)
    build_version_25min(doc)
    build_production_notes(doc)
    doc.save(OUT_FILE)
    print(f"Script vidéo enregistré : {OUT_FILE}")


if __name__ == "__main__":
    build()
