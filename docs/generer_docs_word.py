import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def create_user_manual(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading("Manuel d'Utilisation - Digital School", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Bienvenue dans le manuel d'utilisation de la plateforme Digital School. Ce document est destiné à vous guider dans l'utilisation quotidienne des différentes fonctionnalités du système. Digital School est un ERP complet conçu pour la gestion intégrée des établissements scolaires.")
    
    doc.add_heading("2. Authentification et Accès", level=1)
    doc.add_paragraph("Pour accéder au système, vous devez disposer d'un nom d'utilisateur et d'un mot de passe fournis par l'administrateur.")
    add_bullet(doc, "Rendez-vous sur la page de connexion.")
    add_bullet(doc, "Saisissez vos identifiants.")
    add_bullet(doc, "En cas de perte de mot de passe, contactez l'administrateur système.")
    
    doc.add_heading("3. Module Inscription", level=1)
    doc.add_paragraph("Ce module permet de gérer les admissions et les inscriptions des élèves.")
    add_bullet(doc, "Nouvelle Inscription : Remplissez le formulaire avec les informations de l'élève et de ses parents/tuteurs.")
    add_bullet(doc, "Frais d'inscription : Associez le paiement des frais pour valider l'inscription.")
    add_bullet(doc, "Dossier de l'élève : Consultez et mettez à jour les informations à tout moment.")
    
    doc.add_heading("4. Module Pédagogie", level=1)
    doc.add_paragraph("Ce module est dédié à la gestion des aspects académiques.")
    add_bullet(doc, "Gestion des classes et des emplois du temps.")
    add_bullet(doc, "Saisie et consultation des notes et des bulletins.")
    add_bullet(doc, "Suivi des présences et des absences des élèves.")
    
    doc.add_heading("5. Module Finances", level=1)
    doc.add_paragraph("Le module finance permet le suivi de la trésorerie et la facturation.")
    add_bullet(doc, "Paiements : Enregistrez les paiements de scolarité des élèves.")
    add_bullet(doc, "Dépenses : Suivez les décaissements de l'établissement.")
    add_bullet(doc, "Comptabilité automatique : Génération des écritures comptables (SYSCOHADA) lors de la validation des paiements.")
    
    doc.add_heading("6. Module GRH (Ressources Humaines)", level=1)
    doc.add_paragraph("Gérez le personnel de l'établissement.")
    add_bullet(doc, "Dossiers employés : Informations sur les enseignants et le personnel administratif.")
    add_bullet(doc, "Paie : Gestion des salaires et des avances.")
    add_bullet(doc, "Absences : Suivi des congés et des absences du personnel.")
    
    doc.add_heading("7. Module Utilisateurs et Sécurité", level=1)
    doc.add_paragraph("Gérez les droits d'accès au système.")
    add_bullet(doc, "Création de comptes utilisateurs.")
    add_bullet(doc, "Attribution de rôles (Administrateur, Secrétaire, Comptable, Enseignant).")
    
    doc.save(output_path)
    print(f"User manual saved to {output_path}")

def create_technical_spec(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading("Cahier Technique - Digital School", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    doc.add_heading("1. Présentation de l'Architecture", level=1)
    doc.add_paragraph("Digital School est une application web développée avec le framework Django (Python). L'architecture repose sur le pattern MVT (Model-View-Template), garantissant une séparation claire entre la logique métier, la présentation et la gestion des données.")
    
    doc.add_heading("2. Stack Technologique", level=1)
    add_bullet(doc, "Backend : Python, Django")
    add_bullet(doc, "Base de données : SQLite (environnement de développement) / PostgreSQL (Recommandé pour la production)")
    add_bullet(doc, "Frontend : HTML5, CSS3, JavaScript, Templates Django")
    add_bullet(doc, "Génération de documents : ReportLab (PDF), python-docx (Word)")
    
    doc.add_heading("3. Structure des Modules (Applications Django)", level=1)
    add_bullet(doc, "common : Utilitaires partagés et configurations globales.")
    add_bullet(doc, "utilisateur : Gestion personnalisée des utilisateurs (AbstractUser) et des permissions.")
    add_bullet(doc, "inscription : Modèles de données pour les élèves et les processus d'admission.")
    add_bullet(doc, "pedagogie : Gestion des cours, classes, notes et évaluations.")
    add_bullet(doc, "finances : Suivi des paiements, facturation, et intégration des écritures comptables.")
    add_bullet(doc, "grh : Gestion des ressources humaines, contrats et paie.")
    
    doc.add_heading("4. Modèle de Données (Aperçu)", level=1)
    doc.add_paragraph("Le modèle de données est relationnel. Voici quelques entités principales :")
    add_bullet(doc, "Eleve : Informations personnelles, matricule, classe actuelle.")
    add_bullet(doc, "Paiement : Montant, date, statut (Brouillon, Valide), référence.")
    add_bullet(doc, "Employe : Données RH, type de contrat.")
    
    doc.add_heading("5. Sécurité", level=1)
    doc.add_paragraph("L'application intègre plusieurs mesures de sécurité :")
    add_bullet(doc, "Authentification basée sur les sessions Django.")
    add_bullet(doc, "Protection CSRF sur tous les formulaires (POST).")
    add_bullet(doc, "Contrôle d'accès basé sur les groupes (RBAC) au niveau des vues.")
    
    doc.add_heading("6. Déploiement", level=1)
    doc.add_paragraph("Pour un déploiement en production, il est recommandé d'utiliser Gunicorn comme serveur d'application et Nginx comme reverse proxy. Les fichiers statiques et médias doivent être servis par Nginx.")
    
    doc.save(output_path)
    print(f"Technical spec saved to {output_path}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    manual_path = os.path.join(base_dir, 'Manuel_Utilisation_Digital_School.docx')
    tech_path = os.path.join(base_dir, 'Cahier_Technique_Digital_School.docx')
    
    create_user_manual(manual_path)
    create_technical_spec(tech_path)
