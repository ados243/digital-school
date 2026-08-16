"""
Génère le script de présentation détaillée de Digital School (Word).

Usage :
    python docs/generer_script_presentation.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_FILE = OUT_DIR / "Script_Presentation_Digital_School.docx"


def _set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_p(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run)
    return p


def add_say(doc, text):
    p = doc.add_paragraph()
    label = p.add_run("À dire : ")
    _set_run_font(label, bold=True, color=RGBColor(0, 119, 197))
    body = p.add_run("« " + text + " »")
    _set_run_font(body, italic=True)
    return p


def add_show(doc, text):
    p = doc.add_paragraph()
    label = p.add_run("À montrer : ")
    _set_run_font(label, bold=True, color=RGBColor(22, 101, 52))
    body = p.add_run(text)
    _set_run_font(body)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    label = p.add_run("Note : ")
    _set_run_font(label, bold=True, color=RGBColor(146, 64, 14))
    body = p.add_run(text)
    _set_run_font(body, italic=True)
    return p


def build():
    doc = Document()

    add_title(doc, "Script de présentation — Digital School")
    add_p(
        doc,
        "Document destiné au commercial, au formateur ou au démonstrateur. "
        "Durée recommandée : 45 à 60 minutes (version courte : 25 minutes).",
        italic=True,
    )
    add_p(
        doc,
        "Public cible : direction d’école, promoteur, trésorerie, corps professoral.",
        italic=True,
    )

    add_h(doc, "1. Objectif de la présentation", 1)
    add_bullet(
        doc,
        "Faire comprendre en quoi Digital School remplace les outils dispersés "
        "(cahiers, Excel, WhatsApp informel).",
    )
    add_bullet(
        doc,
        "Démontrer la valeur concrète pour la direction, la caisse, les enseignants, "
        "les parents et les élèves.",
    )
    add_bullet(
        doc,
        "Obtenir un accord de principe pour un essai pilote ou un rendez-vous commercial.",
    )

    add_h(doc, "2. Matériel à préparer", 1)
    add_bullet(
        doc,
        "Comptes démo : administrateur école, caissier, enseignant, parent, élève.",
    )
    add_bullet(
        doc,
        "Données de démo : 1 classe, 3–5 élèves, 1 tuteur, 1 barème Minerval, 1 taux de change.",
    )
    add_bullet(
        doc,
        "Navigateur, partage d’écran, micro ; caméra optionnelle pour la partie visio.",
    )
    add_bullet(
        doc,
        "Connexion internet stable (nécessaire pour Jitsi et WhatsApp).",
    )

    add_h(doc, "3. Structure et timing", 1)
    add_bullet(doc, "Accroche & enjeux — 5 min")
    add_bullet(doc, "Vue d’ensemble de la plateforme — 5 min")
    add_bullet(doc, "Inscriptions & classes — 7 min")
    add_bullet(doc, "Finances & notifications — 10 min")
    add_bullet(doc, "Pédagogie : notes, bulletins, présences — 8 min")
    add_bullet(doc, "Cours en ligne & visioconférence — 8 min")
    add_bullet(doc, "Portails parent / élève / enseignant — 7 min")
    add_bullet(doc, "GRH & paie — 5 min")
    add_bullet(doc, "Conclusion & prochaines étapes — 5 min")
    add_note(
        doc,
        "Version courte (25 min) : accroche, inscriptions, finances, un écran portail parent, conclusion.",
    )

    add_h(doc, "4. Accroche (5 min)", 1)
    add_say(
        doc,
        "Bonjour. Digital School est une plateforme de gestion scolaire conçue pour les écoles "
        "qui veulent centraliser inscriptions, finances, pédagogie et communication parents "
        "dans un seul outil, adapté au contexte congolais : année scolaire nationale, "
        "minerval, CDF/USD, WhatsApp, bulletins TJ et examens.",
    )
    add_say(
        doc,
        "Aujourd’hui, beaucoup d’établissements perdent du temps entre cahiers d’appel, "
        "tableurs de paiements, messages WhatsApp non tracés et bulletins reconstités à la main. "
        "Notre objectif : une information unique, fiable, accessible selon le rôle de chacun.",
    )
    add_show(
        doc,
        "Page de connexion Digital School + logo. Entrer en tant qu’administrateur de l’école démo.",
    )

    add_h(doc, "5. Vue d’ensemble (5 min)", 1)
    add_say(
        doc,
        "Digital School fonctionne en multi-écoles : chaque établissement a son périmètre isolé. "
        "Les rôles sont séparés : direction, caisse, enseignants, parents, élèves. "
        "Chacun voit uniquement ce qui le concerne.",
    )
    add_show(
        doc,
        "Menu latéral : Inscriptions, Pédagogie, Finances, GRH. Pointer les modules sans tout ouvrir.",
    )
    add_bullet(doc, "Inscriptions : élèves, tuteurs, classes, communications.")
    add_bullet(doc, "Finances : frais, encaissements, taux, WhatsApp, comptabilité.")
    add_bullet(doc, "Pédagogie : matières, périodes, classes.")
    add_bullet(doc, "GRH : personnel, contrats, congés, présences, paie.")
    add_bullet(doc, "Portails « Mon espace » : enseignant, parent, élève.")

    add_h(doc, "6. Module Inscriptions (7 min)", 1)
    add_h(doc, "6.1 Classes et structure", 2)
    add_say(
        doc,
        "L’école organise ses classes par section et cycle, avec capacité et titulaire. "
        "L’année scolaire est nationale ; on travaille sur l’année en cours.",
    )
    add_show(
        doc,
        "Inscriptions → Classes & salles : ouvrir une classe, montrer capacité et titulaire.",
    )

    add_h(doc, "6.2 Élèves et parents", 2)
    add_say(
        doc,
        "Chaque élève a une fiche complète, liée à un parent ou tuteur. "
        "Un même tuteur peut être sélectionné pour plusieurs enfants. "
        "La résidence est portée sur la fiche du tuteur.",
    )
    add_show(
        doc,
        "Créer ou ouvrir une fiche élève. Montrer la sélection du tuteur et le popup « Nouveau tuteur ».",
    )
    add_note(doc, "Insister sur le matricule automatique et le rattachement à l’école.")

    add_h(doc, "6.3 Inscription scolaire", 2)
    add_say(
        doc,
        "Inscrire un élève dans une classe pour l’année en cours. "
        "Le système empêche le dépassement de capacité et évite les doubles inscriptions.",
    )
    add_show(
        doc,
        "Nouvelle inscription : élève, classe, année, type. Valider. Revenir à la liste.",
    )

    add_h(doc, "7. Module Finances (10 min)", 1)
    add_h(doc, "7.1 Frais scolaires", 2)
    add_say(
        doc,
        "Vous définissez vos barèmes : minerval, inscription, autres frais. "
        "Le type Minerval est proposé par défaut. "
        "Chaque barème est lié à une section, une année et une échéance.",
    )
    add_show(
        doc,
        "Finances → Frais scolaires → Nouveau. Montrer type de frais (option Ajouter) "
        "et année en cours uniquement.",
    )

    add_h(doc, "7.2 Encaissement", 2)
    add_say(
        doc,
        "Le caissier encaisse en CDF ou USD. Le taux de change est enregistré au moment du paiement. "
        "Un reçu est généré. Les corrections passent par une demande validée par l’admin, "
        "pour garder une piste d’audit.",
    )
    add_show(
        doc,
        "Encaisser un paiement (idéalement avec un compte caissier). Montrer le reçu / l’historique.",
    )

    add_h(doc, "7.3 WhatsApp", 2)
    add_say(
        doc,
        "À chaque paiement validé, le parent peut recevoir une notification WhatsApp "
        "avec les informations du reçu. C’est de la transparence : moins de litiges, plus de confiance.",
    )
    add_show(
        doc,
        "Finances → WhatsApp paiements : config et historique des notifications (même en mode test).",
    )

    add_h(doc, "7.4 Comptabilité", 2)
    add_say(
        doc,
        "Derrière l’encaissement, Digital School alimente la comptabilité de l’école : "
        "plan comptable, journaux, écritures, grand livre et balance. "
        "La direction et la trésorerie ont une vision financière consolidée.",
    )
    add_show(
        doc,
        "Ouvrir rapidement Balance ou Grand livre. Ne pas s’attarder si le public n’est pas comptable.",
    )

    add_h(doc, "8. Pédagogie : notes, bulletins, présences (8 min)", 1)
    add_say(
        doc,
        "Le module pédagogique suit le modèle des bulletins RDC : travaux journaliers et examens "
        "de division, périodes et année en cours. Les enseignants saisissent les notes depuis "
        "leur espace ; le titulaire gère l’appel et les bulletins de classe.",
    )
    add_show(
        doc,
        "Se connecter en enseignant. Montrer : travaux cotés → saisie de notes → aperçu bulletin.",
    )
    add_show(doc, "Présences : appel de classe (compte titulaire).")
    add_note(
        doc,
        "Mentionner l’évolution des élèves : suivi des performances sur les cours du professeur.",
    )

    add_h(doc, "9. Cours en ligne & visioconférence (8 min)", 1)
    add_say(
        doc,
        "Digital School permet aussi de dispenser le cours à distance. "
        "L’enseignant planifie une séance de visioconférence, démarre la salle, "
        "et les élèves rejoignent depuis leur portail. "
        "Pendant le cours, les élèves posent des questions dans un panneau dédié : "
        "l’enseignant répond, épingle ou marque comme traité oralement.",
    )
    add_show(
        doc,
        "Portail enseignant → Cours en ligne → Planifier / Démarrer / Rejoindre la salle.",
    )
    add_show(
        doc,
        "Panneau Questions : poser une question (élève) et répondre (enseignant).",
    )
    add_note(
        doc,
        "Les ressources e-learning (cours enregistrés, chapitres, leçons) restent disponibles "
        "en complément de la visio.",
    )

    add_h(doc, "10. Portails Parent / Élève / Enseignant (7 min)", 1)
    add_h(doc, "10.1 Parent", 2)
    add_say(
        doc,
        "Le parent se connecte à Mon espace, voit ses enfants, consulte notes, présences, "
        "situation des frais, lit les annonces de la direction et écrit au titulaire.",
    )
    add_show(
        doc,
        "Compte parent : liste des enfants → détail d’un enfant → annonces / messages.",
    )

    add_h(doc, "10.2 Élève", 2)
    add_say(
        doc,
        "L’élève voit sa situation scolaire, ses dernières notes, peut étudier les ressources "
        "et rejoindre les cours en visioconférence.",
    )
    add_show(doc, "Compte élève : Mon espace → Cours en ligne.")

    add_h(doc, "10.3 Enseignant", 2)
    add_say(
        doc,
        "L’enseignant a un cockpit : classes, travaux, notes, présences (si titulaire), "
        "cours en ligne, messagerie parents.",
    )
    add_show(doc, "Dashboard enseignant : carte d’une classe.")

    add_h(doc, "11. GRH & paie (5 min)", 1)
    add_say(
        doc,
        "La gestion du personnel complète le dispositif : fiches, contrats, congés, "
        "pointage, génération de paie. Les salaires peuvent être payés depuis le module Finances.",
    )
    add_show(doc, "GRH → Personnel / Paies. Finances → Payer le personnel (aperçu).")

    add_h(doc, "12. Points différenciants à marteler", 1)
    add_bullet(doc, "Tout-en-un : administration, caisse, pédagogie, RH, portails.")
    add_bullet(
        doc,
        "Pensé pour le terrain RDC : minerval, CDF/USD, WhatsApp, bulletins TJ/examen.",
    )
    add_bullet(doc, "Rôles et isolation multi-écoles.")
    add_bullet(doc, "Traçabilité des paiements et des corrections.")
    add_bullet(doc, "Cours à distance avec visio + questions en direct.")
    add_bullet(doc, "Communication direction ↔ parents intégrée.")

    add_h(doc, "13. Objections fréquentes & réponses", 1)
    add_bullet(
        doc,
        "« On n’a pas de connexion stable. » → Le back-office tourne sur navigateur ; "
        "la visio nécessite internet. On démarre par inscriptions + finances, la visio en phase 2.",
    )
    add_bullet(
        doc,
        "« Nos enseignants ne sont pas à l’aise avec l’informatique. » → Interfaces portail "
        "simplifiées, formation courte, rôles limités à l’essentiel.",
    )
    add_bullet(
        doc,
        "« On a déjà Excel / un logiciel. » → Digital School relie caisse, pédagogie et parents : "
        "moins de ressaisie, moins d’erreurs, plus de contrôle.",
    )
    add_bullet(
        doc,
        "« Combien ça coûte ? » → Offre adaptée à la taille de l’école (licence, hébergement, "
        "formation, support). Proposer un pilote avant engagement long.",
    )

    add_h(doc, "14. Conclusion & appel à l’action (5 min)", 1)
    add_say(
        doc,
        "En résumé, Digital School donne à votre établissement une colonne vertébrale numérique : "
        "des élèves bien inscrits, une caisse maîtrisée, une pédagogie suivie, "
        "des parents informés et des enseignants équipés pour enseigner aussi à distance.",
    )
    add_say(
        doc,
        "Je vous propose trois suites possibles : une démo approfondie sur vos données, "
        "un essai pilote de quelques semaines, ou une offre commerciale adaptée à la taille "
        "de votre école. Quelle option vous convient le mieux ?",
    )
    add_show(
        doc,
        "Laisser l’écran sur le dashboard. Remettre la lettre de proposition / plaquette.",
    )

    add_h(doc, "15. Checklist post-présentation", 1)
    add_bullet(
        doc,
        "Noter les décideurs présents et leurs priorités (caisse, bulletins, parents…).",
    )
    add_bullet(doc, "Envoyer le récap + lien de démo sous 24 h.")
    add_bullet(doc, "Proposer une date d’essai pilote.")
    add_bullet(
        doc,
        "Identifier le référent technique côté école (connexion, WhatsApp, comptes).",
    )

    add_h(doc, "Annexe A — Parcours démo minute par minute (version 45 min)", 1)
    rows = [
        ("0–5", "Accroche + connexion admin"),
        ("5–10", "Tour du menu / rôles"),
        ("10–17", "Élève + tuteur + inscription"),
        ("17–27", "Frais → paiement → WhatsApp / reçu"),
        ("27–35", "Notes / bulletin / présence (enseignant)"),
        ("35–42", "Visio + questions + portail parent"),
        ("42–45", "Conclusion & prochaines étapes"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Temps"
    hdr[1].text = "Étape"
    for t, step in rows:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = step

    add_h(doc, "Annexe B — Phrases clés (mémo)", 1)
    add_bullet(doc, "Une seule source de vérité pour l’école.")
    add_bullet(doc, "Chaque rôle voit juste ce qu’il doit voir.")
    add_bullet(
        doc,
        "Le parent est informé ; la caisse est traçable ; le bulletin est structuré.",
    )
    add_bullet(doc, "Le cours continue même à distance, avec interaction.")

    add_h(doc, "Annexe C — Ordre de connexion recommandé pendant la démo", 1)
    add_bullet(doc, "1. Admin / direction — vue d’ensemble + inscriptions + frais")
    add_bullet(doc, "2. Caissier — encaissement")
    add_bullet(doc, "3. Enseignant — notes, présence, cours visio")
    add_bullet(doc, "4. Élève — rejoindre la salle / poser une question")
    add_bullet(doc, "5. Parent — suivi enfant, frais, annonces")

    doc.save(OUT_FILE)
    print(f"Script enregistré : {OUT_FILE}")


if __name__ == "__main__":
    build()
