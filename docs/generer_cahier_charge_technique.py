"""
Génère le Cahier des charges technique + Catalogue fonctionnel Digital School
avec logos NTT S.A.R.L et Digital School.

Usage :
    python docs/generer_cahier_charge_technique.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm, Twips


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Cahier_Charges_Technique_Digital_School.docx"
LOGO_NTT = ROOT / "assets" / "logo-ntt.png"
LOGO_DS = ROOT / "assets" / "logo-ds.png"

BLUE = RGBColor(0, 119, 197)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
WHITE = RGBColor(255, 255, 255)


def set_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), kwargs[edge].get("val", "single"))
            tag.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            tag.set(qn("w:color"), kwargs[edge].get("color", "0077C5"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=BLUE if level > 1 else DARK)
    return h


def add_p(doc, text, size=11, bold=False, italic=False, color=None, center=False, space_after=8):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color or DARK)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True, color=DARK)
        r2 = p.add_run(text)
        set_run(r2, color=DARK)
    else:
        r = p.add_run(text)
        set_run(r, color=DARK)
    return p


def feature_table(doc, rows):
    """rows: list of (fonctionnalité, description)"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = ""
    hdr[0].paragraphs[0].clear()
    hdr[1].paragraphs[0].clear()
    r0 = hdr[0].paragraphs[0].add_run("Fonctionnalité")
    r1 = hdr[1].paragraphs[0].add_run("Description détaillée")
    set_run(r0, bold=True, color=WHITE, size=10)
    set_run(r1, bold=True, color=WHITE, size=10)
    shade_cell(hdr[0], "0077C5")
    shade_cell(hdr[1], "0077C5")
    hdr[0].width = Cm(5.5)
    hdr[1].width = Cm(11.5)

    for i, (feat, desc) in enumerate(rows):
        row = table.add_row().cells
        row[0].paragraphs[0].clear()
        row[1].paragraphs[0].clear()
        a = row[0].paragraphs[0].add_run(feat)
        b = row[1].paragraphs[0].add_run(desc)
        set_run(a, bold=True, size=10, color=DARK)
        set_run(b, size=10, color=DARK)
        if i % 2 == 1:
            shade_cell(row[0], "E8F4FD")
            shade_cell(row[1], "E8F4FD")
    doc.add_paragraph()


def cover_page(doc):
    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Logos row
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1, c2 = table.rows[0].cells
    c0.width = Cm(6)
    c1.width = Cm(5)
    c2.width = Cm(6)

    if LOGO_NTT.is_file():
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_NTT), width=Inches(1.6))
        t = c0.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("NTT S.A.R.L")
        set_run(r, size=10, bold=True, color=DARK)
        t2 = c0.add_paragraph()
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t2.add_run("Ntobu's Technology")
        set_run(r2, size=8, italic=True, color=MUTED)

    mid = c1.paragraphs[0]
    mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = mid.add_run("×")
    set_run(rm, size=28, bold=True, color=BLUE)

    if LOGO_DS.is_file():
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_DS), width=Inches(1.5))
        t = c2.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("Digital School")
        set_run(r, size=10, bold=True, color=BLUE)

    doc.add_paragraph()
    doc.add_paragraph()

    # Blue banner title
    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, "0077C5")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CAHIER DES CHARGES TECHNIQUE")
    set_run(r, size=20, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("& CATALOGUE FONCTIONNEL COMPLET")
    set_run(r2, size=14, bold=True, color=WHITE)

    doc.add_paragraph()
    add_p(doc, "Plateforme de gestion scolaire intégrée", size=16, bold=True, center=True, color=DARK)
    add_p(
        doc,
        "Document de spécification technique et présentation détaillée de l’ensemble des fonctionnalités",
        size=11,
        italic=True,
        center=True,
        color=MUTED,
    )
    doc.add_paragraph()
    add_p(doc, "Éditeur : NTT S.A.R.L — Ntobu's Technology", size=12, center=True, bold=True)
    add_p(doc, "Produit : Digital School", size=12, center=True, color=BLUE, bold=True)
    add_p(doc, "Version document : 1.0  ·  Août 2026", size=10, center=True, color=MUTED)
    add_p(doc, "Classification : Confidentiel — usage commercial & déploiement", size=9, italic=True, center=True, color=MUTED)

    doc.add_page_break()


def intro_ntt(doc):
    add_h(doc, "1. Présentation de l’éditeur — NTT S.A.R.L", 1)
    if LOGO_NTT.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO_NTT), width=Inches(1.2))

    add_p(
        doc,
        "NTT S.A.R.L (Ntobu’s Technology) est une startup technologique spécialisée dans la conception "
        "de solutions numériques pour les organisations. Digital School est son produit phare dédié "
        "à la modernisation de la gestion des établissements scolaires.",
    )
    add_bullet(doc, "Raison sociale : NTT S.A.R.L (Ntobu’s Technology)")
    add_bullet(doc, "Activité : édition de logiciels, digitalisation, solutions métiers")
    add_bullet(doc, "Produit : Digital School — ERP scolaire multi-établissements")
    add_bullet(doc, "Positionnement : outils concrets pour le terrain éducatif (RDC / Afrique)")
    add_p(
        doc,
        "Le présent document constitue le cahier des charges technique de référence et le catalogue "
        "exhaustif des fonctionnalités livrées dans la plateforme.",
        italic=True,
        color=MUTED,
    )


def intro_product(doc):
    add_h(doc, "2. Présentation du produit — Digital School", 1)
    if LOGO_DS.is_file():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(LOGO_DS), width=Inches(1.1))

    add_p(
        doc,
        "Digital School est une plateforme web de gestion scolaire intégrée. Elle centralise "
        "les inscriptions, les finances (minerval, caisse, comptabilité), la pédagogie "
        "(notes, bulletins RDC, présences), les ressources humaines, les portails parent / élève / enseignant, "
        "ainsi que la dispensation de cours en ligne et en visioconférence.",
    )
    add_h(doc, "2.1 Objectifs", 2)
    add_bullet(doc, "Remplacer les outils dispersés (cahiers, Excel, WhatsApp informel).")
    add_bullet(doc, "Garantir une source unique de vérité par établissement.")
    add_bullet(doc, "Donner à chaque rôle (direction, caisse, enseignant, parent, élève) un espace adapté.")
    add_bullet(doc, "Assurer la traçabilité financière et pédagogique.")
    add_bullet(doc, "Permettre l’enseignement à distance (visio + questions + ressources).")

    add_h(doc, "2.2 Publics cibles", 2)
    add_bullet(doc, "Écoles privées, publiques et conventionnées")
    add_bullet(doc, "Directions, trésorerie / caisses, corps professoral")
    add_bullet(doc, "Parents / tuteurs et élèves")


def tech_specs(doc):
    add_h(doc, "3. Cahier des charges technique", 1)

    add_h(doc, "3.1 Architecture générale", 2)
    add_p(
        doc,
        "Application web monolithique structurée en modules Django (pattern MVT), "
        "avec isolation multi-établissements par colonne ecole_id (multi-tenant logique) "
        "et portails utilisateurs séparés du back-office.",
    )
    add_bullet(doc, "Backend : Python 3 · Django 4.1 (MVT)")
    add_bullet(doc, "Base de données : MySQL (production recommandée) / configurable via .env")
    add_bullet(doc, "Front-end : templates Django, HTML5, CSS3 (charte bleue #0077C5), JavaScript")
    add_bullet(doc, "Authentification : modèle Utilisateur custom (AbstractUser)")
    add_bullet(doc, "Fichiers médias : photos, avatars, documents de cours")
    add_bullet(doc, "Exports : Excel (openpyxl) pour situations de paiements")

    add_h(doc, "3.2 Modules applicatifs (apps)", 2)
    feature_table(
        doc,
        [
            ("inscription", "Écoles, élèves, tuteurs, classes, inscriptions, années scolaires"),
            ("finances", "Frais, paiements, taux, WhatsApp, salaires, comptabilité OHADA"),
            ("pedagogie", "Matières, périodes bulletin, travaux, notes, présences, e-learning, visio"),
            ("grh", "Personnel, contrats, congés, pointage, paies"),
            ("utilisateur", "Comptes, rôles, portails, messagerie, profil, reset MDP"),
        ],
    )

    add_h(doc, "3.3 Rôles et sécurité", 2)
    feature_table(
        doc,
        [
            ("Manager / Directeur / Trésorier", "Back-office complet (inscriptions, finances, pédagogie, GRH)"),
            ("Caissier", "Espace restreint : encaissements élèves + profil"),
            ("Professeur / Enseignant", "Portail Mon espace (classes, notes, cours, visio, messages)"),
            ("Parent", "Portail enfants, frais, notes, annonces, messagerie"),
            ("Élève", "Portail notes + cours en ligne / visioconférence"),
            ("Superutilisateur", "Administration Django + configuration référentiels"),
        ],
    )
    add_p(
        doc,
        "Un middleware de restriction (RestrictionPortailMiddleware) confine chaque profil "
        "à ses URL autorisées. Les données métier sont filtrées systématiquement par école.",
        italic=True,
        color=MUTED,
    )

    add_h(doc, "3.4 Intégrations externes", 2)
    add_bullet(doc, "WhatsApp Business : Ultramsg, Meta Cloud API, ou mode test (LOG)")
    add_bullet(doc, "Visioconférence : Jitsi Meet (domaine configurable JITSI_DOMAIN)")
    add_bullet(doc, "E-mail SMTP : récupération de mot de passe (Brevo / SMTP école en production)")

    add_h(doc, "3.5 Exigences non fonctionnelles", 2)
    add_bullet(doc, "Disponibilité : hébergement HTTPS, sauvegardes BDD quotidiennes recommandées")
    add_bullet(doc, "Performance : Gunicorn/Nginx en production ; médias hors disque local (S3/Spaces) recommandé")
    add_bullet(doc, "Traçabilité : reçus numérotés, demandes de correction, journaux WhatsApp, écritures comptables")
    add_bullet(doc, "Localisation : interface française ; devises CDF/USD ; bulletins modèle RDC")
    add_bullet(doc, "Évolutivité : multi-écoles sur une même instance")

    add_h(doc, "3.6 Environnement de déploiement recommandé", 2)
    add_bullet(doc, "Serveur Linux (VPS) : Ubuntu LTS")
    add_bullet(doc, "Processus applicatif : Gunicorn + Nginx + Certificat SSL")
    add_bullet(doc, "Variables d’environnement : .env (SECRET_KEY, DB_*, EMAIL_*, JITSI_*, WhatsApp)")
    add_bullet(doc, "Commandes : migrate, collectstatic, création superuser / seed référentiels")


def features_catalog(doc):
    add_h(doc, "4. Catalogue complet des fonctionnalités", 1)
    add_p(
        doc,
        "Cette section énumère en détail l’ensemble des fonctionnalités user-facing livrées "
        "dans Digital School, module par module.",
        italic=True,
        color=MUTED,
    )

    # --- Auth ---
    add_h(doc, "4.1 Authentification, comptes et profil", 2)
    feature_table(
        doc,
        [
            ("Connexion", "Authentifie l’utilisateur via identifiant et mot de passe."),
            ("Déconnexion", "Termine la session de manière sécurisée."),
            ("Création de compte", "Auto-inscription Parent, Élève ou Corps professoral liée au matricule école."),
            ("Redirection post-login", "Oriente vers le portail, l’espace enseignant ou le back-office selon le rôle."),
            ("Mot de passe oublié", "Envoie un lien de réinitialisation par e-mail (identifiant ou e-mail)."),
            ("Nouveau mot de passe", "Permet de définir un nouveau mot de passe via lien sécurisé (token)."),
            ("Mon profil", "Modifie prénom, nom, e-mail, téléphone, photo de profil."),
            ("Changement de mot de passe", "Met à jour le mot de passe lorsque l’utilisateur est connecté."),
            ("Restriction par rôle", "Limite l’accès aux menus et URL selon le profil (middleware)."),
        ],
    )

    # --- Inscriptions ---
    add_h(doc, "4.2 Module Inscriptions & Élèves", 2)
    feature_table(
        doc,
        [
            ("Tableau de bord inscriptions", "Synthétise effectifs, répartition F/G, capacités et derniers enregistrements."),
            ("Liste / recherche élèves", "Consulte les fiches élèves avec recherche et statut d’inscription année en cours."),
            ("Création / modification / suppression élève", "Gère l’identité, nationalité, photo et rattachement tuteur (matricule ELV-…)."),
            ("Gestion des tuteurs / parents", "Enregistre coordonnées, résidence, lien de parenté (matricule TUT-…) ; popup AJAX depuis la fiche élève."),
            ("Liste globale des parents", "Permet de sélectionner un tuteur déjà existant dans toute l’application."),
            ("Inscriptions scolaires", "Inscrit un élève dans une classe pour l’année en cours (nouvelle, réinscription, transfert)."),
            ("Numérotation automatique", "Attribue des numéros d’inscription INS-AAAA-AAAA-####."),
            ("Contrôle de capacité", "Empêche le dépassement de la capacité maximale d’une classe."),
            ("Badge « Déjà inscrit »", "Indique si l’élève est déjà inscrit pour l’année scolaire active."),
            ("Classes & salles", "Gère section, salle, capacité, titulaire et taux de remplissage."),
            ("Communications direction", "Envoie des annonces ciblées (école, section, classe, parent)."),
            ("Référentiels", "Écoles, cycles, sections, communes/quartiers, années scolaires nationales (admin)."),
        ],
    )

    # --- Finances ---
    add_h(doc, "4.3 Module Finances & Paiements", 2)
    feature_table(
        doc,
        [
            ("Dashboard finances", "Affiche encaissements, caisse, KPI et synthèse minerval par classe."),
            ("Types de frais", "Catégorise les frais ; type Minerval créé par défaut pour chaque école."),
            ("Ajout de type depuis le formulaire", "Option « Ajouter » + popup pour créer un type sans quitter le barème."),
            ("Frais scolaires (barèmes)", "Définit montant, devise, section, échéance, caractère obligatoire, année en cours."),
            ("Encaissement paiement", "Saisit un paiement (espèces, Mobile Money, virement, chèque) avec reçu."),
            ("Multi-devises CDF / USD", "Enregistre le montant reçu et convertit via le taux de change du jour."),
            ("Taux de change", "Historise le taux (CDF pour 1 USD) par établissement."),
            ("Impression de reçu", "Produit le reçu numéroté du paiement."),
            ("Demandes de correction", "Le caissier signale une erreur ; l’admin valide ou rejette la correction."),
            ("Paiements par classe", "Situation de classe exportable en Excel."),
            ("Inscriptions non payées", "Liste les dossiers dont les frais d’inscription restent dus."),
            ("Paiement des salaires", "Règle les paies du personnel (unitaire ou par lot) depuis les finances."),
            ("Espace caissier dédié", "Interface restreinte centrée sur l’encaissement."),
        ],
    )

    add_h(doc, "4.4 WhatsApp paiements", 2)
    feature_table(
        doc,
        [
            ("Configuration par école", "Paramètre Ultramsg, Meta Cloud API ou mode test."),
            ("Modèles Meta & variables", "Mappe les variables du template Business approuvé."),
            ("Notification automatique", "Envoie un message au parent à chaque paiement validé."),
            ("Test d’envoi", "Vérifie la configuration depuis l’écran WhatsApp."),
            ("Renvoi de notification", "Relance un envoi pour un paiement donné."),
            ("Journal des notifications", "Trace destinataire, statut et réponse API."),
        ],
    )

    add_h(doc, "4.5 Comptabilité (référentiel OHADA / HOADA)", 2)
    feature_table(
        doc,
        [
            ("Plan comptable", "Gère les comptes (numéro, libellé, devise) par école."),
            ("Journaux comptables", "Organise les journaux (caisse, banque, etc.)."),
            ("Écritures", "Saisit des écritures multi-lignes débit/crédit."),
            ("Grand livre", "Consulte le grand livre des comptes."),
            ("Balance", "Produit la balance des comptes."),
            ("Posting automatique paiements", "Génère les écritures à la validation d’un encaissement."),
            ("Posting automatique salaires", "Génère les écritures lors du paiement d’une paie."),
            ("Resynchronisation", "Met à jour la compta après correction ou annulation de paiement."),
        ],
    )

    # --- Pedagogie ---
    add_h(doc, "4.6 Module Pédagogie", 2)
    feature_table(
        doc,
        [
            ("Dashboard pédagogie", "Vue de synthèse du module académique."),
            ("Catalogue de matières", "Code, libellé, coefficient, maxima bulletin RDC, section, enseignant de référence."),
            ("Périodes de bulletin", "Trimestres / semestres et périodes d’évaluation selon le cycle (calendrier RDC)."),
            ("Activation période / division", "Marque la période ou division « en cours » pour la saisie."),
            ("Affectations enseignant", "Attribue un professeur à une matière pour une classe."),
            ("Travaux cotés", "Crée devoirs, interros, examens, TP (TJ ou examen de division)."),
            ("Saisie des notes", "Encode les notes élève par travail, avec règles d’absence / verrouillage."),
            ("Bulletins scolaires RDC", "Calcule TJ, examens, totaux et pourcentages (classe / élève)."),
            ("Évolution des élèves", "Visualise la progression des notes pour le professeur de cours."),
            ("Appel de présence", "Saisit présent / absent / retard / excusé (titulaire de classe)."),
            ("Récapitulatif d’assiduité", "Bilan de présence par élève sur l’année."),
        ],
    )

    # --- Elearning / Visio ---
    add_h(doc, "4.7 Cours en ligne & visioconférence", 2)
    add_h(doc, "4.7.1 Ressources asynchrones (e-learning)", 3)
    feature_table(
        doc,
        [
            ("Création de cours", "Compose un parcours (objectifs, niveau, couverture, classe/matière)."),
            ("Chapitres & leçons", "Structure le contenu (vidéo, lecture, document, exercice)."),
            ("Publication contrôlée", "Rend le cours visible aux élèves de la classe."),
            ("Bibliothèque élève", "Liste les cours publiés de sa classe / année."),
            ("Progression", "Suit les leçons vues et terminées."),
        ],
    )
    add_h(doc, "4.7.2 Cours en direct (Jitsi)", 3)
    feature_table(
        doc,
        [
            ("Planification de séance", "Définit titre, classe, matière, date/heure, durée."),
            ("Démarrer / terminer / annuler", "Pilote le cycle de vie de la visioconférence."),
            ("Salle enseignant / élève", "Rejoint la salle Jitsi embarquée (ou nouvel onglet)."),
            ("Questions en direct", "Les élèves posent des questions pendant le cours."),
            ("Réponses & épinglage", "L’enseignant répond, épingle ou marque « traité oralement »."),
            ("Rafraîchissement live", "Le panneau questions se met à jour automatiquement."),
        ],
    )

    # --- Portails ---
    add_h(doc, "4.8 Portails utilisateurs (Mon espace)", 2)
    add_h(doc, "4.8.1 Enseignant", 3)
    feature_table(
        doc,
        [
            ("Dashboard classes", "Vue des classes titulaire / cours, effectifs et matières."),
            ("Fiche classe", "Élèves, tuteurs, travaux, présences et bulletins (si titulaire)."),
            ("Travaux & notes", "Gestion complète des évaluations depuis le portail."),
            ("Messagerie parents", "Contacte le tuteur d’un élève de la classe."),
            ("Cours en ligne / visio", "Publie des ressources et anime des séances live."),
        ],
    )
    add_h(doc, "4.8.2 Élève", 3)
    feature_table(
        doc,
        [
            ("Mon espace", "Affiche classe, année et dernières notes."),
            ("Cours en ligne", "Accède aux ressources et rejoint les visioconférences."),
        ],
    )
    add_h(doc, "4.8.3 Parent", 3)
    feature_table(
        doc,
        [
            ("Mes enfants", "Liste les enfants rattachés via la fiche tuteur."),
            ("Suivi détaillé", "Notes, présences, situation des frais, titulaire."),
            ("Annonces direction", "Lit les communications officielles."),
            ("Messagerie", "Échange avec le professeur titulaire."),
        ],
    )

    # --- GRH ---
    add_h(doc, "4.9 Module GRH (Ressources humaines)", 2)
    feature_table(
        doc,
        [
            ("Dashboard GRH", "Synthèse RH de l’établissement."),
            ("Fiches personnel", "Matricules PER-…, fonctions, contacts, photo, lien compte utilisateur."),
            ("Contrats", "CDI, CDD, prestataire, stage — salaire de base, devise, statut."),
            ("Congés", "Demandes (annuel, maladie, etc.) avec approbation / rejet."),
            ("Présences personnel", "Pointage présent / absent / retard / congé."),
            ("Paies", "Génération mensuelle (base, primes, déductions, net) et marquage payé."),
        ],
    )

    # --- Multi-tenant ---
    add_h(doc, "4.10 Multi-établissements", 2)
    feature_table(
        doc,
        [
            ("Entité École", "Fiche établissement (code, type, contacts, activation)."),
            ("Isolation des données", "Toutes les données métier filtrées par école connectée."),
            ("Matricules indépendants", "Séquences ELV / TUT / PER / reçus propres à chaque école."),
            ("Configs par école", "WhatsApp, taux de change, plan comptable, types de frais."),
            ("Année scolaire nationale", "Calendrier partagé (MINEDU), une année « en cours »."),
        ],
    )


def matrix_roles(doc):
    add_h(doc, "5. Matrice des accès par rôle", 1)
    rows = [
        ("Fonctionnalité", "Direction", "Caissier", "Enseignant", "Parent", "Élève"),
        ("Inscriptions / classes", "Oui", "Non", "Lecture portail", "Non", "Non"),
        ("Encaissements", "Oui", "Oui", "Non", "Non", "Non"),
        ("WhatsApp / Comptabilité", "Oui", "Limité", "Non", "Non", "Non"),
        ("Notes / bulletins", "Oui", "Non", "Oui", "Consultation", "Consultation"),
        ("Présences élèves", "Oui", "Non", "Titulaire", "Consultation", "Non"),
        ("Cours visio", "Non*", "Non", "Oui", "Non", "Oui"),
        ("Messagerie / annonces", "Oui", "Non", "Oui", "Oui", "Non"),
        ("GRH / paie", "Oui", "Non", "Non", "Non", "Non"),
        ("Mon profil", "Oui", "Oui", "Oui", "Oui", "Oui"),
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, h in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=9, color=WHITE)
        shade_cell(cell, "0077C5")
    for line in rows[1:]:
        row = table.add_row().cells
        for i, val in enumerate(line):
            row[i].paragraphs[0].clear()
            r = row[i].paragraphs[0].add_run(val)
            set_run(r, size=9, bold=(i == 0), color=DARK)
    add_p(doc, "* Accès possible via comptes enseignant dédiés.", size=9, italic=True, color=MUTED)


def closing(doc):
    add_h(doc, "6. Livrables et prochaines étapes", 1)
    add_bullet(doc, "Application Digital School opérationnelle (modules listés ci-dessus)")
    add_bullet(doc, "Documentation : manuels, scripts de présentation, présent document")
    add_bullet(doc, "Déploiement : VPS / cloud + configuration .env production")
    add_bullet(doc, "Formation direction, caisse, enseignants, communication parents")
    add_bullet(doc, "Essai pilote recommandé avant généralisation multi-écoles")

    doc.add_paragraph()
    # Footer brand table
    t = doc.add_table(rows=1, cols=2)
    c0, c1 = t.rows[0].cells
    shade_cell(c0, "0F172A")
    shade_cell(c1, "0077C5")
    c0.paragraphs[0].clear()
    c1.paragraphs[0].clear()
    r0 = c0.paragraphs[0].add_run("NTT S.A.R.L — Ntobu's Technology")
    set_run(r0, bold=True, color=WHITE, size=11)
    r1 = c1.paragraphs[0].add_run("Digital School — Plateforme scolaire intégrée")
    set_run(r1, bold=True, color=WHITE, size=11)

    add_p(
        doc,
        "Document généré pour présentation commerciale, réponse à appel d’offres et cadrage de déploiement.",
        size=9,
        italic=True,
        center=True,
        color=MUTED,
    )


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    cover_page(doc)
    intro_ntt(doc)
    intro_product(doc)
    tech_specs(doc)
    features_catalog(doc)
    matrix_roles(doc)
    closing(doc)

    doc.save(OUT)
    print(f"Document généré : {OUT}")


if __name__ == "__main__":
    build()
