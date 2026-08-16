"""
Génère la proposition financière Digital School destinée à une école.
Modèle : 20 USD par carte d'élève = droit d'utiliser l'application.

Usage :
    python docs/generer_proposition_financiere.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Proposition_Financiere_Digital_School.docx"
LOGO_NTT = ROOT / "assets" / "logo-ntt.png"
LOGO_DS = ROOT / "assets" / "logo-ds.png"

BLUE = RGBColor(0, 119, 197)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
WHITE = RGBColor(255, 255, 255)
GREEN = RGBColor(22, 101, 52)

PRICE = 20  # USD par carte d'élève / année scolaire
PRICE_DUP = 10  # USD duplicata


def fmt_amount(n):
    return f"{int(n):,}".replace(",", " ")


def fmt_dec(n):
    return f"{n:.2f}".replace(".", ",")


def set_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run(
            run,
            size={1: 16, 2: 13, 3: 12}.get(level, 11),
            bold=True,
            color=BLUE if level > 1 else DARK,
        )
    return h


def add_p(doc, text, size=11, bold=False, italic=False, color=None, center=False, space_after=8):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color or DARK)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True, color=DARK)
        r2 = p.add_run(text)
        set_run(r2, color=DARK)
    else:
        r = p.add_run(text)
        set_run(r, color=DARK)
    return p


def cell_text(cell, text, bold=False, size=10, color=None, center=False, white=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    set_run(r, bold=bold, size=size, color=WHITE if white else (color or DARK))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def styled_table(doc, headers, rows, col_widths=None, highlight_last=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_text(cell, h, bold=True, size=10, white=True, center=True)
        shade_cell(cell, "0077C5")
    prevent_row_split(table.rows[0])

    for ri, line in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[ri + 1])
        last = highlight_last and ri == len(rows) - 1
        for i, val in enumerate(line):
            cell_text(
                cells[i],
                val,
                bold=(i == 0 or last),
                size=10,
                center=(i > 0),
                color=GREEN if last and i == len(line) - 1 else DARK,
            )
            if last:
                shade_cell(cells[i], "DCFCE7")
            elif ri % 2 == 1:
                shade_cell(cells[i], "E8F4FD")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def price_banner(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "0077C5")
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{PRICE} USD  ·  par élève  ·  par année scolaire")
    set_run(r, size=18, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Une carte d’élève = l’accès complet à Digital School pour l’établissement")
    set_run(r2, size=11, italic=True, color=WHITE)
    doc.add_paragraph()


def highlight_box(doc, title, body, fill="E8F4FD"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run(r2, size=11, color=DARK)
    doc.add_paragraph()


def cover_page(doc):
    for _ in range(1):
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1, c2 = table.rows[0].cells
    c0.width = Cm(6)
    c1.width = Cm(5)
    c2.width = Cm(6)

    if LOGO_NTT.is_file():
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_NTT), width=Inches(1.5))
        t = c0.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("NTT S.A.R.L")
        set_run(r, size=10, bold=True, color=DARK)
        t2 = c0.add_paragraph()
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t2.add_run("Ntobu's Technology")
        set_run(r2, size=8, italic=True, color=MUTED)

    mid = c1.paragraphs[0]
    mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = mid.add_run("×")
    set_run(rm, size=28, bold=True, color=BLUE)

    if LOGO_DS.is_file():
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_DS), width=Inches(1.4))
        t = c2.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("Digital School")
        set_run(r, size=10, bold=True, color=BLUE)

    doc.add_paragraph()
    doc.add_paragraph()

    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, "0077C5")
    set_cell_margins(cell, top=160, bottom=160)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROPOSITION FINANCIÈRE")
    set_run(r, size=22, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Offre école  —  Cartes d’élève & accès à la plateforme")
    set_run(r2, size=13, bold=True, color=WHITE)

    doc.add_paragraph()
    add_p(doc, "Digital School", size=18, bold=True, center=True, color=BLUE)
    add_p(
        doc,
        "Plateforme de gestion scolaire intégrée",
        size=12,
        italic=True,
        center=True,
        color=MUTED,
    )
    doc.add_paragraph()

    meta = [
        ("Éditeur", "NTT S.A.R.L — Ntobu's Technology"),
        ("Produit", "Digital School"),
        ("Destinataire", "Direction de l’établissement  [Nom de l’école]"),
        ("Référence", "PROP-DS-2026-001"),
        ("Date", "Août 2026"),
        ("Validité de l’offre", "90 jours à compter de la date du présent document"),
        ("Devise", "USD (dollars américains)"),
        ("Classification", "Confidentiel — usage commercial"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        c_k, c_v = t.rows[i].cells
        c_k.width = Cm(5)
        c_v.width = Cm(11.5)
        cell_text(c_k, k, bold=True, size=10, white=True)
        shade_cell(c_k, "0F172A" if i % 2 == 0 else "0077C5")
        cell_text(c_v, v, size=10)
        shade_cell(c_v, "F8FAFC")

    doc.add_paragraph()
    add_p(
        doc,
        "Document destiné à la direction, au promoteur et à la trésorerie de l’établissement.",
        size=9,
        italic=True,
        center=True,
        color=MUTED,
    )
    doc.add_page_break()


def objet(doc):
    add_h(doc, "1. Objet de la proposition", 1)
    add_p(
        doc,
        "La présente proposition financière est adressée à votre établissement. "
        "Elle décrit les conditions commerciales d’accès à Digital School, "
        "la plateforme de gestion scolaire conçue par NTT S.A.R.L.",
    )
    add_p(
        doc,
        "Le principe est volontairement simple : l’école ne paie pas de licence logicielle, "
        "ni d’abonnement mensuel. Elle commande des cartes d’élève. "
        "Chaque carte ouvre le droit d’utiliser l’application pour cet élève, "
        "et donne à l’établissement l’accès à l’ensemble des modules.",
    )
    price_banner(doc)


def principe(doc):
    add_h(doc, "2. Principe de l’offre", 1)
    add_h(doc, "2.1 Une carte, un élève, une année", 2)
    add_p(
        doc,
        "Pour chaque élève inscrit, l’école commande une carte Digital School au tarif unique "
        f"de {PRICE} USD. Cette carte vaut pour l’année scolaire en cours.",
    )
    add_bullet(doc, "1 carte = 1 élève = 1 année scolaire.")
    add_bullet(doc, "Pas de frais d’installation, pas de licence, pas d’abonnement caché.")
    add_bullet(doc, "Les comptes direction, caisse, enseignants et parents sont illimités et inclus.")
    add_bullet(doc, "Tous les modules de Digital School sont ouverts dès la première commande.")
    add_bullet(doc, "Les nouveaux élèves en cours d’année font l’objet d’une commande complémentaire au même tarif.")

    add_h(doc, "2.2 Pourquoi ce modèle convient aux écoles", 2)
    add_p(
        doc,
        "Les établissements facturent déjà une carte scolaire aux parents. "
        "Avec Digital School, cette carte devient à la fois badge d’identification "
        "et droit d’accès à un outil complet de gestion.",
    )
    highlight_box(
        doc,
        "Coût net possible pour l’école : 0 USD",
        f"Si l’établissement répercute les {PRICE} USD aux parents au titre de la carte d’élève "
        "(pratique déjà courante), l’école n’engage aucun budget informatique. "
        "Elle obtient les cartes physiques et l’usage de toute la plateforme.",
        fill="DCFCE7",
    )

    add_h(doc, "2.3 Équivalent journalier", 2)
    add_p(
        doc,
        f"{PRICE} USD sur une année scolaire, c’est environ {fmt_dec(PRICE / 12)} USD par mois et par élève, "
        f"soit moins de {fmt_dec(PRICE / 270)} USD par jour. Un montant inférieur au prix d’un transport "
        "ou d’une photocopie de devoir, pour un outil utilisé toute l’année "
        "par la direction, la caisse, les enseignants et les parents.",
    )


def inclus(doc):
    add_h(doc, f"3. Ce qui est inclus dans les {PRICE} USD", 1)
    add_p(
        doc,
        f"Le tarif de {PRICE} USD par élève couvre à la fois le support physique "
        "et le droit d’usage logiciel. Rien n’est facturé à part pour « ouvrir » un module.",
    )

    add_h(doc, "3.1 La carte d’élève", 2)
    add_bullet(doc, "Carte nominative de l’année scolaire (identité, classe, matricule).")
    add_bullet(doc, "Photo de l’élève, nom de l’établissement, année en cours.")
    add_bullet(doc, "Code visuel (QR) rattaché au dossier Digital School.")
    add_bullet(doc, "Support durable, adapté au port quotidien (contrôle d’accès, examens, caisse).")

    add_h(doc, "3.2 La plateforme Digital School — tous les modules", 2)
    styled_table(
        doc,
        ["Module", "Ce que l’école peut faire"],
        [
            (
                "Inscriptions",
                "Élèves, tuteurs, classes, sections, cycles, inscriptions annuelles, capacité des salles.",
            ),
            (
                "Finances",
                "Frais et minerval, encaissements CDF/USD, taux de change, reçus, WhatsApp, comptabilité OHADA.",
            ),
            (
                "Pédagogie",
                "Matières, notes, travaux, bulletins modèle RDC, présences, cours en ligne et visioconférence.",
            ),
            (
                "GRH & paie",
                "Personnel, contrats, congés, pointage, salaires, fiches de paie.",
            ),
            (
                "Portails",
                "Espaces direction, caisse, enseignant, parent et élève — chacun voit ce qui le concerne.",
            ),
            (
                "Communication",
                "Messagerie, annonces par classe ou par école, suivi de lecture.",
            ),
        ],
        col_widths=[4.5, 12.5],
    )

    add_h(doc, "3.3 Services d’accompagnement inclus", 2)
    add_bullet(doc, "Mise en service de l’école sur la plateforme (compte établissement, rôles, année scolaire).")
    add_bullet(doc, "Hébergement, sauvegardes et mises à jour de Digital School pendant l’année couverte par les cartes.")
    add_bullet(doc, "Formation initiale : 1 session pour la direction / la caisse, 1 session pour les enseignants.")
    add_bullet(doc, "Support de base (e-mail / WhatsApp) pendant toute l’année scolaire.")
    add_bullet(doc, "Comptes utilisateurs illimités (personnel, parents, élèves porteurs de carte).")


def simulations(doc):
    add_h(doc, "4. Simulations financières", 1)
    add_p(
        doc,
        f"Le montant dû est uniquement le nombre d’élèves × {PRICE} USD. "
        "Les tableaux ci-dessous permettent à la direction de situer l’engagement "
        "selon la taille réelle de l’établissement.",
    )

    effectifs = [50, 100, 150, 200, 300, 400, 500, 800, 1000, 1500]
    mois = fmt_dec(PRICE / 12)
    rows = [
        (
            f"{fmt_amount(n)} élèves",
            fmt_amount(n),
            fmt_amount(n * PRICE),
            f"{mois} USD",
        )
        for n in effectifs
    ]

    add_h(doc, "4.1 Barème selon l’effectif", 2)
    styled_table(
        doc,
        ["Effectif élèves", "Cartes à commander", "Montant année (USD)", "Soit / mois / élève"],
        rows,
        col_widths=[4.2, 4.2, 4.5, 4.1],
    )
    add_p(
        doc,
        f"Formule : montant = nombre d’élèves inscrits × {PRICE} USD. "
        "Aucun palier, aucune surprise au-delà de ce calcul.",
        size=10,
        italic=True,
        color=MUTED,
    )

    exemple_n = 400
    exemple_total = fmt_amount(exemple_n * PRICE)
    add_h(doc, f"4.2 Exemple chiffré — école de {fmt_amount(exemple_n)} élèves", 2)
    styled_table(
        doc,
        ["Poste", "Détail", "Montant"],
        [
            ("Commande de cartes", f"{fmt_amount(exemple_n)} élèves × {PRICE} USD", f"{exemple_total} USD"),
            ("Licence logiciel", "Incluse", "0 USD"),
            ("Hébergement & mises à jour", "Inclus pour l’année", "0 USD"),
            ("Comptes enseignants / parents", "Illimités", "0 USD"),
            ("Formation initiale", "2 sessions incluses", "0 USD"),
            ("Total dû à NTT S.A.R.L", "Pour l’année scolaire", f"{exemple_total} USD"),
            (f"Si répercuté aux parents ({PRICE} USD/carte)", f"{fmt_amount(exemple_n)} × {PRICE} USD encaissés", f"{exemple_total} USD"),
            ("Coût net pour l’école", "Cartes + logiciel complets", "0 USD"),
        ],
        col_widths=[6.0, 6.5, 4.5],
        highlight_last=True,
    )

    add_h(doc, "4.3 Lecture pour le promoteur et la trésorerie", 2)
    add_bullet(doc, "Le prix suit l’effectif réel : une petite école paie moins qu’une grande.")
    add_bullet(doc, "Pas d’engagement pluriannuel obligatoire : on recommande chaque année scolaire.")
    add_bullet(
        doc,
        f"La trésorerie peut intégrer les {PRICE} USD dans les frais d’inscription ou les facturer comme « carte scolaire ».",
    )
    add_bullet(
        doc,
        f"Les élèves arrivés en cours d’année s’ajoutent au prorata des cartes commandées (tarif plein {PRICE} USD, pas de découpe mensuelle).",
    )


def hors_offre(doc):
    add_h(doc, "5. Prestations hors offre de base", 1)
    add_p(
        doc,
        "Tout ce qui suit est facultatif. L’école peut fonctionner sans ces lignes. "
        "Elles ne conditionnent pas l’accès à Digital School.",
    )
    styled_table(
        doc,
        ["Prestation", "Tarif indicatif", "Quand y recourir"],
        [
            ("Duplicata de carte (perte, détérioration)", f"{PRICE_DUP} USD / carte", "Élève déjà inscrit, carte à réimprimer."),
            ("Cartes supplémentaires en cours d’année", f"{PRICE} USD / élève", "Nouvelles inscriptions après la commande initiale."),
            ("Import de données historiques (Excel)", "Sur devis", "Reprise d’un ancien fichier élèves / paiements."),
            ("Formation supplémentaire sur site", "50 USD / session", "Au-delà des 2 sessions incluses."),
            ("Développement sur mesure", "Sur devis", "Module spécifique demandé par l’école."),
            ("Matériel informatique", "Non fourni", "Ordinateurs, imprimantes, connexion internet restent à charge de l’école."),
        ],
        col_widths=[6.5, 4.5, 6.0],
    )
    add_p(
        doc,
        "Un élève qui quitte l’établissement en cours d’année n’ouvre pas droit à remboursement : "
        "la carte a été produite et le droit d’usage ouvert pour l’année.",
        size=10,
        italic=True,
        color=MUTED,
    )


def deroulement(doc):
    add_h(doc, "6. Déroulement de la collaboration", 1)
    styled_table(
        doc,
        ["Étape", "Qui fait quoi", "Délai indicatif"],
        [
            ("1. Accord", "La direction valide la présente proposition (effectif estimé).", "Jour J"),
            ("2. Fiche établissement", "NTT crée l’école, l’année scolaire et les comptes direction / caisse.", "48 h"),
            ("3. Liste des élèves", "L’école transmet la liste (nom, prénom, classe, photo si disponible).", "Selon l’école"),
            ("4. Commande & paiement", "Commande ferme du nombre de cartes. Paiement selon les modalités ci-dessous.", "À la commande"),
            ("5. Production des cartes", "NTT fabrique et livre les cartes nominatives.", "7 à 15 jours ouvrés"),
            ("6. Activation", "Les élèves porteurs de carte sont actifs dans Digital School.", "Dès production"),
            ("7. Formation", "Sessions direction / caisse puis enseignants.", "Dans les 15 jours"),
            ("8. Année suivante", "Nouvelle commande selon l’effectif de la rentrée.", "Avant la rentrée"),
        ],
        col_widths=[4.0, 9.0, 4.0],
    )

    add_h(doc, "6.1 Modalités de paiement", 2)
    add_bullet(doc, "Devise : USD.")
    add_bullet(doc, "Acompte de 50 % à la commande ferme, solde à la livraison des cartes.")
    add_bullet(doc, "Pour une première commande inférieure à 100 cartes : paiement intégral à la commande.")
    add_bullet(doc, "Moyens acceptés : virement bancaire, Mobile Money (M-Pesa, Airtel Money, Orange Money), espèces contre reçu.")
    add_bullet(doc, "Une facture / note de débit NTT S.A.R.L est émise à chaque commande.")
    add_bullet(doc, "Merci d’indiquer en communication : nom de l’école + réf. PROP-DS-2026-001.")

    add_h(doc, "6.2 Coordonnées bancaires — NTOBO’S TECHNOLOGY SARL", 2)
    add_p(
        doc,
        "Les virements sont à effectuer exclusivement sur l’un des comptes ci-dessous, "
        "ouverts au nom de NTOBO’S TECHNOLOGY / NTOBO’S TECHNOLOGY SARL.",
    )
    styled_table(
        doc,
        ["Banque", "N° de compte", "Intitulé"],
        [
            ("UBA RDC", "030320004374", "NTOBO’S TECHNOLOGY"),
            ("EQUITY BCDC", "00011150722200037532797", "NTOBO’S TECHNOLOGY SARL"),
            ("ECOBANK RDC", "35900000801", "NTOBO’S TECHNOLOGY SARL"),
        ],
        col_widths=[4.0, 7.0, 6.0],
    )
    add_p(
        doc,
        "Tout paiement doit être suivi d’une preuve (bordereau, capture Mobile Money ou reçu) "
        "adressée à NTT S.A.R.L pour imputation de la commande.",
        size=10,
        italic=True,
        color=MUTED,
    )

    add_h(doc, "6.3 Durée et renouvellement", 2)
    add_p(
        doc,
        "L’accès à Digital School est ouvert pour l’année scolaire correspondant aux cartes commandées. "
        "Au renouvellement, l’école commande les cartes de la nouvelle année selon l’effectif réel. "
        "Les données de l’école (dossiers, historiques de paiements, notes) sont conservées : "
        "on ne recommence pas de zéro.",
    )


def comparaison(doc):
    add_h(doc, "7. Ce que l’école gagne concrètement", 1)
    add_p(
        doc,
        "Digital School remplace les outils dispersés (cahiers d’appel, Excel de caisse, "
        "WhatsApp non tracé, bulletins reconstitués à la main) par une seule source de vérité.",
    )
    styled_table(
        doc,
        ["Aujourd’hui, sans plateforme", f"Avec Digital School (via la carte à {PRICE} USD)"],
        [
            ("Cahiers, fichiers Excel, doubles saisies", "Dossier unique élève / tuteur / classe"),
            ("Caisse difficile à contrôler", "Reçus numérotés, CDF/USD, traçabilité, OHADA"),
            ("Parents relancés au cas par cas", "Portail parent + notifications WhatsApp"),
            ("Bulletins longs à produire", "Bulletins modèle RDC calculés par le système"),
            ("Enseignants isolés", "Notes, présences, cours en ligne, visio"),
            ("Licence logicielle + maintenance à part", "Inclus dans la carte — rien d’autre à budgéter"),
        ],
        col_widths=[8.5, 8.5],
    )
    highlight_box(
        doc,
        "Message à retenir pour le conseil de gestion",
        "Vous n’achetez pas un logiciel. Vous commandez les cartes d’élève de l’année. "
        "L’application Digital School — inscriptions, caisse, pédagogie, GRH, portails — "
        "est le service qui accompagne ces cartes.",
    )


def conditions(doc):
    add_h(doc, "8. Conditions commerciales", 1)
    add_bullet(doc, "Offre valable 90 jours à compter de la date figurant en page de garde.")
    add_bullet(doc, f"Tarif unique : {PRICE} USD par carte d’élève, par année scolaire, hors duplicata ({PRICE_DUP} USD).")
    add_bullet(doc, "Commande minimale conseillée : 30 cartes (en dessous, un devis spécifique peut être établi).")
    add_bullet(doc, "Les photos et listes d’élèves restent la propriété de l’établissement ; NTT les utilise uniquement pour produire les cartes et alimenter Digital School.")
    add_bullet(doc, "L’école s’engage à commander une carte pour chaque élève inscrit utilisant la plateforme.")
    add_bullet(doc, "NTT s’engage à maintenir la plateforme disponible pendant l’année couverte, hors cas de force majeure.")
    add_bullet(doc, "Les présentes conditions prévalent sur tout échange oral antérieur. Un bon de commande signé vaut acceptation.")
    add_p(
        doc,
        "Le présent document constitue une proposition commerciale et non une facture. "
        "L’engagement naît à la signature du bon de commande et au versement de l’acompte.",
        italic=True,
        color=MUTED,
    )


def bon_commande(doc):
    add_h(doc, "9. Bon de commande (à compléter)", 1)
    add_p(
        doc,
        "Merci de photocopier ou d’imprimer cette page, de la compléter et de la renvoyer "
        "à NTT S.A.R.L. Un devis / facture conforme vous sera adressé sous 48 heures.",
        italic=True,
        color=MUTED,
    )

    fields = [
        ("Nom de l’établissement", ""),
        ("Ville / commune", ""),
        ("Nom du promoteur / directeur", ""),
        ("Téléphone", ""),
        ("E-mail", ""),
        ("Année scolaire concernée", "20____ / 20____"),
        ("Nombre de cartes commandées", f"________  élèves  ×  {PRICE} USD"),
        ("Montant total (USD)", "________"),
        ("Duplicatas éventuels", f"________  ×  {PRICE_DUP} USD  =  ________"),
        ("Montant général", "________  USD"),
        ("Mode de paiement souhaité", "Virement  /  Mobile Money  /  Espèces"),
        ("Banque de virement", "UBA RDC  /  EQUITY BCDC  /  ECOBANK RDC"),
    ]
    t = doc.add_table(rows=len(fields), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, hint) in enumerate(fields):
        a, b = t.rows[i].cells
        a.width = Cm(6.5)
        b.width = Cm(10.5)
        prevent_row_split(t.rows[i])
        cell_text(a, label, bold=True, size=10, white=True)
        shade_cell(a, "0077C5")
        cell_text(b, hint if hint else " ", size=10)
        shade_cell(b, "F8FAFC")
        set_cell_margins(b, top=90, bottom=90)

    doc.add_paragraph()
    add_p(
        doc,
        "Je soussigné(e), dûment habilité(e) à engager l’établissement, accepte la proposition "
        "financière Digital School (réf. PROP-DS-2026-001) et commande le nombre de cartes indiqué ci-dessus.",
        size=10,
    )

    sig = doc.add_table(rows=2, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_sig = ("Pour l’établissement", "Pour NTT S.A.R.L")
    bodies = (
        "Nom :\nFonction :\nDate :\nSignature & cachet :",
        "Nom :\nFonction :\nDate :\nSignature & cachet :",
    )
    for i in range(2):
        cell_text(sig.rows[0].cells[i], headers_sig[i], bold=True, size=11, white=True, center=True)
        shade_cell(sig.rows[0].cells[i], "0F172A")
        sig.rows[1].cells[i].paragraphs[0].clear()
        p = sig.rows[1].cells[i].paragraphs[0]
        r = p.add_run(bodies[i])
        set_run(r, size=10, color=DARK)
        set_cell_margins(sig.rows[1].cells[i], top=120, bottom=200, left=120, right=120)
        shade_cell(sig.rows[1].cells[i], "F8FAFC")
        sig.rows[1].cells[i].width = Cm(8.5)

    doc.add_paragraph()


def closing(doc):
    add_h(doc, "10. Contact", 1)
    add_p(
        doc,
        "Pour une démonstration, un ajustement d’effectif ou toute question sur cette offre :",
    )
    add_bullet(doc, "Éditeur : NTT S.A.R.L — Ntobu's Technology")
    add_bullet(doc, "Produit : Digital School")
    add_bullet(doc, "E-mail : [e-mail commercial]")
    add_bullet(doc, "Téléphone / WhatsApp : [numéro]")
    add_bullet(doc, "Référence à rappeler : PROP-DS-2026-001")

    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    c0, c1 = t.rows[0].cells
    shade_cell(c0, "0F172A")
    shade_cell(c1, "0077C5")
    set_cell_margins(c0, top=80, bottom=80)
    set_cell_margins(c1, top=80, bottom=80)
    c0.paragraphs[0].clear()
    c1.paragraphs[0].clear()
    r0 = c0.paragraphs[0].add_run("NTT S.A.R.L — Ntobu's Technology")
    set_run(r0, bold=True, color=WHITE, size=11)
    r1 = c1.paragraphs[0].add_run(f"Digital School — {PRICE} USD par carte d’élève")
    set_run(r1, bold=True, color=WHITE, size=11)

    add_p(
        doc,
        "Proposition financière — Août 2026 — Document confidentiel, non contractuel tant que le bon de commande n’est pas signé.",
        size=9,
        italic=True,
        center=True,
        color=MUTED,
    )


def set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NTT S.A.R.L  ·  Digital School  ·  Proposition financière PROP-DS-2026-001  ·  Confidentiel")
    set_run(r, size=8, color=MUTED, italic=True)

    # Page number field
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Page ")
    set_run(r2, size=8, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r3 = p2.add_run()
    r3._r.append(fld1)
    r3._r.append(instr)
    r3._r.append(fld2)
    set_run(r3, size=8, color=MUTED)
    r4 = p2.add_run(" / ")
    set_run(r4, size=8, color=MUTED)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    r5 = p2.add_run()
    r5._r.append(fld3)
    r5._r.append(instr2)
    r5._r.append(fld4)
    set_run(r5, size=8, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    cover_page(doc)
    objet(doc)
    principe(doc)
    inclus(doc)
    simulations(doc)
    hors_offre(doc)
    deroulement(doc)
    comparaison(doc)
    conditions(doc)
    bon_commande(doc)
    closing(doc)
    set_footer(doc)

    try:
        doc.save(OUT)
        print(f"Document généré : {OUT}")
    except PermissionError:
        alt = ROOT / "Proposition_Financiere_Digital_School_maj.docx"
        doc.save(alt)
        print(f"Fichier original ouvert (verrouillé). Document généré : {alt}")


if __name__ == "__main__":
    build()
